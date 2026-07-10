from __future__ import annotations

from pathlib import Path
import builtins

import pytest

from backend.app.assignments.parser import parse_assignment_document


def test_assignment_txt_parsing(tmp_path: Path):
    path = tmp_path / "brief.txt"
    path.write_text("Big Data Portfolio\nBuild a Kafka pipeline.", encoding="utf-8")

    parsed = parse_assignment_document(path)

    assert parsed.title == "Big Data Portfolio"
    assert "Kafka pipeline" in parsed.extracted_text
    assert parsed.source_path == str(path)
    assert parsed.document_id.startswith("assignment-")


def test_assignment_markdown_parsing(tmp_path: Path):
    path = tmp_path / "brief.md"
    path.write_text("# Portfolio Brief\n\nAssignment 1: Kafka + Grafana", encoding="utf-8")

    parsed = parse_assignment_document(path)

    assert parsed.title == "Portfolio Brief"
    assert "Assignment 1" in parsed.extracted_text


def test_assignment_docx_parsing_if_dependency_available(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "brief.docx"
    document = Document()
    document.add_paragraph("Docx Assignment Brief")
    document.add_paragraph("Take screenshots of the dashboard.")
    document.save(path)

    parsed = parse_assignment_document(path)

    assert parsed.title == "Docx Assignment Brief"
    assert "screenshots" in parsed.extracted_text


def test_assignment_docx_missing_dependency_is_controlled(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    path = tmp_path / "brief.docx"
    path.write_bytes(b"not opened when import fails")
    original_import = builtins.__import__

    def fake_import(name, globals=None, locals=None, fromlist=(), level=0):
        if name == "docx":
            raise ImportError("missing docx")
        return original_import(name, globals, locals, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    with pytest.raises(ValueError, match="python-docx is required"):
        parse_assignment_document(path)


def test_assignment_folder_path_returns_clear_error(tmp_path: Path):
    with pytest.raises(ValueError, match="folder"):
        parse_assignment_document(tmp_path)


def test_assignment_missing_file_handling(tmp_path: Path):
    with pytest.raises(FileNotFoundError):
        parse_assignment_document(tmp_path / "missing.txt")


def test_assignment_unsupported_extension_handling(tmp_path: Path):
    path = tmp_path / "brief.pdf"
    path.write_text("not supported", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported"):
        parse_assignment_document(path)
