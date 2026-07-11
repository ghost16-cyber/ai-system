from __future__ import annotations

import hashlib
import json
import subprocess
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from backend.app.main import create_app


ASSIGNMENT = "assignment-2"
WORKSPACE = "assignment_workspaces/assignment_2"


def _setup(tmp_path: Path) -> tuple[TestClient, Path]:
    workspace = tmp_path / WORKSPACE
    workspace.mkdir(parents=True)
    client = TestClient(create_app(tmp_path / "app.db", workspace_root=tmp_path))
    return client, workspace


def _requirements() -> list[dict]:
    return [
        {"requirement_id": "code", "title": "Python implementation", "description": "Create the implementation", "source_reference": "task-1", "requirement_category": "code_file", "required_deliverable_type": "python", "expected_evidence": ["main.py"], "verification_method": "file"},
        {"requirement_id": "screenshot", "title": "Dashboard screenshot", "description": "Include dashboard evidence", "source_reference": "task-2", "requirement_category": "screenshot", "required_deliverable_type": "image", "expected_evidence": ["dashboard.png"], "verification_method": "manual"},
        {"requirement_id": "missing", "title": "Observed discussion", "description": "Discuss observed results", "source_reference": "task-3", "requirement_category": "report", "required_deliverable_type": "report", "expected_evidence": ["discussion.md"], "verification_method": "manual"},
    ]


def _verify(client: TestClient):
    response = client.post(f"/assignments/{ASSIGNMENT}/verify", json={"workspace_path": WORKSPACE, "assignment_output": {"requirements": _requirements()}})
    assert response.status_code == 200, response.text
    return response.json()


def _create(client: TestClient, title: str | None = None):
    response = client.post(f"/assignments/{ASSIGNMENT}/reports", json={"workspace_path": WORKSPACE, "title": title})
    assert response.status_code == 200, response.text
    return response.json()


def test_report_creation_is_snapshot_bound_and_uses_placeholders(tmp_path: Path) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('grounded implementation')\n", encoding="utf-8")
    (workspace / "dashboard.png").write_bytes(b"dashboard screenshot evidence bytes")
    with client:
        snapshot = _verify(client)
        report = _create(client)
        cross = client.get(f"/assignments/assignment-3/reports/{report['report_id']}", params={"workspace_path": WORKSPACE})
    assert report["source_verification_snapshot"] == snapshot["snapshot_id"]
    assert report["workspace"] == WORKSPACE
    states = {section["section_id"]: section["verification_state"] for section in report["report_sections"]}
    assert states["implementation"] == "partially_supported"
    assert states["evidence"] == "requires_manual_review"
    assert report["unresolved_items"]
    assert cross.status_code == 404


def test_manual_acceptance_included_and_rejected_evidence_excluded(tmp_path: Path) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('implementation evidence')\n", encoding="utf-8")
    (workspace / "dashboard.png").write_bytes(b"dashboard screenshot evidence bytes")
    with client:
        _verify(client)
        accepted = client.post(f"/assignments/{ASSIGNMENT}/evidence/review", json={"workspace_path": WORKSPACE, "requirement_id": "screenshot", "evidence_reference": "file:dashboard.png", "decision": "accepted", "note": "Reviewed"})
        assert accepted.status_code == 200
        client.post(f"/assignments/{ASSIGNMENT}/verify", json={"workspace_path": WORKSPACE, "assignment_output": {"requirements": _requirements()}})
        accepted_report = _create(client)
        rejected = client.post(f"/assignments/{ASSIGNMENT}/evidence/review", json={"workspace_path": WORKSPACE, "requirement_id": "screenshot", "evidence_reference": "file:dashboard.png", "decision": "rejected", "note": "Unclear"})
        assert rejected.status_code == 200
        client.post(f"/assignments/{ASSIGNMENT}/verify", json={"workspace_path": WORKSPACE, "assignment_output": {"requirements": _requirements()}})
        rejected_report = _create(client)
    accepted_section = next(item for item in accepted_report["report_sections"] if item["section_id"] == "evidence")
    rejected_section = next(item for item in rejected_report["report_sections"] if item["section_id"] == "evidence")
    assert accepted_section["verification_state"] == "manually_accepted"
    assert "file:dashboard.png" in accepted_section["linked_evidence"]
    assert "file:dashboard.png" not in rejected_section["linked_evidence"]
    assert "Rejected evidence" in " ".join(rejected_section["warnings"])


def test_stale_conflicting_and_failed_states_propagate_without_favorable_choice(tmp_path: Path) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('implementation evidence')\n", encoding="utf-8")
    (workspace / "dashboard.png").write_bytes(b"dashboard screenshot evidence bytes")
    with client:
        snapshot = _verify(client)
        snapshot_path = next((tmp_path / "data" / "assignment_verification" / "snapshots").glob("*.json"))
        payload = json.loads(snapshot_path.read_text(encoding="utf-8"))
        payload["requirements"][0]["status"] = "partially_verified"
        payload["requirements"][0]["warnings"] = ["Execution evidence is stale because files changed."]
        payload["requirements"][0]["linked_execution_evidence"] = ["assignment-command:stale"]
        payload["requirements"][1]["warnings"] = ["Conflicting passing and failed evidence exists."]
        snapshot_path.write_text(json.dumps(payload), encoding="utf-8")
        report = _create(client)
        readiness = client.get(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/readiness", params={"workspace_path": WORKSPACE}).json()
    implementation = next(item for item in report["report_sections"] if item["section_id"] == "implementation")
    evidence = next(item for item in report["report_sections"] if item["section_id"] == "evidence")
    assert implementation["verification_state"] == "stale"
    assert evidence["verification_state"] == "requires_manual_review"
    assert readiness["stale_evidence_count"] >= 1
    assert readiness["export_blockers"]
    assert snapshot["assignment_id"] == ASSIGNMENT


def test_revision_history_safe_edits_and_immutable_evidence(tmp_path: Path) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('implementation evidence')\n", encoding="utf-8")
    with client:
        _verify(client)
        snapshot_path = next((tmp_path / "data" / "assignment_verification" / "snapshots").glob("*.json"))
        before = hashlib.sha256(snapshot_path.read_bytes()).hexdigest()
        report = _create(client)
        optional = next(item for item in report["report_sections"] if not item["mandatory"])
        updated = client.patch(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}", json={"workspace_path": WORKSPACE, "changes": {"title": "User report password=hidden", "sections": [{"section_id": optional["section_id"], "user_editable_notes": "My observation", "inclusion_status": "excluded"}]}})
        forbidden = client.patch(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}", json={"workspace_path": WORKSPACE, "changes": {"sections": [{"section_id": "objectives", "inclusion_status": "excluded"}]}})
    body = updated.json()
    assert updated.status_code == 200
    assert len(body["revisions"]) == 2
    assert body["revisions"][-1]["previous_revision_reference"] == report["current_revision_id"]
    assert "hidden" not in body["title"]
    assert forbidden.status_code == 400
    assert hashlib.sha256(snapshot_path.read_bytes()).hexdigest() == before


def test_markdown_json_docx_and_zip_exports_are_safe_and_selected_only(tmp_path: Path) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('implementation evidence')\n", encoding="utf-8")
    (workspace / "notes.txt").write_text("not selected\n", encoding="utf-8")
    with client:
        _verify(client); report = _create(client, "Safe Report")
        records = {}
        for fmt in ("markdown", "json", "docx", "zip"):
            response = client.post(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/export", json={"workspace_path": WORKSPACE, "format": fmt, "selected_files": ["main.py"] if fmt == "zip" else []})
            assert response.status_code == 200, response.text
            records[fmt] = response.json()
        exports = client.get(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/exports", params={"workspace_path": WORKSPACE}).json()
        downloads = {fmt: client.get(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/exports/{record['export_id']}", params={"workspace_path": WORKSPACE}) for fmt, record in records.items()}
    assert len(exports["exports"]) == 4
    assert downloads["markdown"].headers["content-type"].startswith("text/markdown")
    assert downloads["json"].json()["absolute_paths_included"] is False
    assert downloads["docx"].content.startswith(b"PK")
    from docx import Document
    document = Document(BytesIO(downloads["docx"].content))
    assert document.sections[0].top_margin.inches == 1.0
    assert any(paragraph.text == "Requirement Traceability" for paragraph in document.paragraphs)
    assert document.tables and len(document.tables[-1].columns) == 3
    zip_path = tmp_path / "package.zip"; zip_path.write_bytes(downloads["zip"].content)
    with zipfile.ZipFile(zip_path) as archive:
        names = archive.namelist()
        assert "workspace/main.py" in names
        assert "workspace/notes.txt" not in names
        manifest = json.loads(archive.read("report/traceability-manifest.json"))
    selected = manifest["selected_files"][0]
    assert selected["sha256"] == hashlib.sha256((workspace / "main.py").read_bytes()).hexdigest()
    assert all(not Path(record["filename"]).is_absolute() for record in records.values())


def test_package_path_symlink_absolute_and_dangerous_types_rejected(tmp_path: Path) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('implementation evidence')\n", encoding="utf-8")
    outside = tmp_path / "outside.py"; outside.write_text("outside\n", encoding="utf-8")
    (workspace / "linked.py").symlink_to(outside)
    (workspace / "danger.exe").write_bytes(b"danger")
    with client:
        # Remove the symlink before verification because inventory correctly rejects it.
        (workspace / "linked.py").unlink(); _verify(client); report = _create(client); (workspace / "linked.py").symlink_to(outside)
        payload = lambda selected: {"workspace_path": WORKSPACE, "format": "zip", "selected_files": selected}
        traversal = client.post(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/export", json=payload(["../outside.py"]))
        absolute = client.post(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/export", json=payload([str(outside)]))
        symlink = client.post(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/export", json=payload(["linked.py"]))
        dangerous = client.post(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/export", json=payload(["danger.exe"]))
    assert {traversal.status_code, absolute.status_code, symlink.status_code, dangerous.status_code} == {400}


def test_corrupt_and_unsupported_report_records_fail_safely(tmp_path: Path) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('implementation evidence')\n", encoding="utf-8")
    with client:
        _verify(client); report = _create(client)
        record_path = next((tmp_path / "data" / "assignment_reports").rglob(f"{report['report_id']}.json"))
        payload = json.loads(record_path.read_text(encoding="utf-8")); payload["schema_version"] = 999; record_path.write_text(json.dumps(payload), encoding="utf-8")
        unsupported = client.get(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}", params={"workspace_path": WORKSPACE})
        record_path.write_text("{corrupt", encoding="utf-8")
        corrupt = client.get(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}", params={"workspace_path": WORKSPACE})
    assert unsupported.status_code == 400
    assert "version" in unsupported.json()["detail"].lower()
    assert corrupt.status_code == 400


def test_assembly_and_export_never_invoke_subprocess_or_guarantee_correctness(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    client, workspace = _setup(tmp_path)
    (workspace / "main.py").write_text("def main():\n    print('implementation evidence')\n", encoding="utf-8")
    monkeypatch.setattr(subprocess, "Popen", lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("subprocess forbidden")))
    with client:
        _verify(client); report = _create(client)
        exported = client.post(f"/assignments/{ASSIGNMENT}/reports/{report['report_id']}/export", json={"workspace_path": WORKSPACE, "format": "markdown", "selected_files": []})
    assert exported.status_code == 200
    assert "guarantee" not in json.dumps(report).lower()
