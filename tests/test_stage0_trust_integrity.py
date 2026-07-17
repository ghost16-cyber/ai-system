from __future__ import annotations

import copy
import hashlib
import sqlite3
from pathlib import Path

import pytest

from backend.app.assignments.code_blueprints import generate_code_blueprints
from backend.app.assignments.dataset_mapper import map_dataset_columns
from backend.app.assignments.extractor import extract_assignment_brief
from backend.app.assignments.parser import (
    AssignmentDocumentLimits,
    DocumentLimitExceeded,
    parse_assignment_document,
)
from backend.app.datasets.schemas import DatasetProfile, DatasetSuitability
from backend.app.assignments.schemas import DatasetSemanticMapping, DerivedColumnPlan, SourceColumn
from backend.app.database.repository import AnalysisRepository
from backend.app.project_analysis.state_manifest import (
    IncompleteProjectManifestError,
    ProjectManifestLimits,
    assert_manifest_fresh,
    build_project_state_manifest,
)
from backend.app.project_delivery import (
    ProjectDeliveryError,
    ProjectVerifierError,
    VerificationMode,
    VerificationState,
    activate_next_work_unit,
    approve_plan,
    assert_verifier_result_fresh,
    create_delivery_job,
    generate_handoff,
    record_verification,
    run_deterministic_verifier,
)


def _delivery_project(root: Path) -> Path:
    project = root / "project"
    project.mkdir()
    (project / "README.md").write_text("Feature: greet returns Hello, Ada!\n", encoding="utf-8")
    (project / "app.py").write_text("def greet(name):\n    return f'Hello, {name}!'\n", encoding="utf-8")
    (project / "test_app.py").write_text(
        "from app import greet\n\ndef test_greet():\n    assert greet('Ada') == 'Hello, Ada!'\n",
        encoding="utf-8",
    )
    (project / "pyproject.toml").write_text(
        "[tool.pytest.ini_options]\ntestpaths=['.']\n", encoding="utf-8"
    )
    (project / "unrelated.py").write_text("VALUE = 'protected'\n", encoding="utf-8")
    return project


def _delivery_job(project: Path) -> dict:
    return create_delivery_job(
        root=project,
        conversation_id="conversation",
        folder_access_id="workspace",
        user_request="Deliver the project change in README.md by implementing app.py.",
        action_run_id="run",
    )


def _household_profile(path: str = "household_power_consumption.csv") -> DatasetProfile:
    columns = [
        "Date", "Time", "Global_active_power", "Global_reactive_power",
        "Voltage", "Global_intensity", "Sub_metering_1",
    ]
    return DatasetProfile(
        dataset_path=path, detected_format="csv", detected_delimiter=";",
        row_count_estimate=2_075_259, column_count=len(columns), columns=columns,
        detected_date_columns=["Date"],
        detected_numeric_columns=columns[2:], detected_categorical_columns=[],
        sample_rows_limited=[{
            "Date": "16/12/2006", "Time": "17:24:00", "Global_active_power": "4.216",
        }],
        suitability=DatasetSuitability(
            assignment_1_suitable=True, assignment_2_suitable=True,
            assignment_3_suitable=True, reasons=["Timestamp and numeric measures are available."],
            recommended_assignment_use="Use the same source for all three assignments.",
        ),
    )


def test_docx_blocks_preserve_mixed_body_order_and_provenance(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "three-assignments.docx"
    document = Document()
    document.add_heading("ASSIGNMENT ONE — Kafka and Grafana (30 Marks)", level=1)
    document.add_paragraph("Task 1: Build a Kafka producer and capture screenshots.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Task"
    table.cell(0, 1).text = "Requirement"
    table.cell(1, 0).text = "Task 2"
    table.cell(1, 1).text = "Create a Grafana dashboard (10 marks)."
    document.add_heading("Assignment 2: Spark and Snowflake — 30 Marks", level=2)
    document.add_paragraph("Task 1: Aggregate the supplied dataset.")
    document.add_paragraph("ASSIGNMENT THREE – Streaming and Redis (30 Marks)", style="Heading 3")
    document.add_paragraph("Task 1: Build the live dashboard.")
    document.save(path)

    parsed = parse_assignment_document(path)
    visible = [
        block.text for block in parsed.document_blocks
        if block.block_type in {"heading", "paragraph", "table_row"}
    ]
    assert visible.index("Task 1: Build a Kafka producer and capture screenshots.") < visible.index("Task | Requirement")
    assert visible.index("Task 2 | Create a Grafana dashboard (10 marks).") < visible.index("Assignment 2: Spark and Snowflake — 30 Marks")
    assert [block.order_index for block in parsed.document_blocks] == list(range(len(parsed.document_blocks)))
    assert all(block.source_span.source_id for block in parsed.document_blocks)
    assert all(block.block_id.startswith("doc-block-") for block in parsed.document_blocks)

    brief = extract_assignment_brief(parsed)
    assert [section.title.split()[1].rstrip(":") for section in brief.sections] == ["ONE", "2", "THREE"]
    assert len(brief.sections) == 3
    assert all(section.source_spans and section.source_block_ids for section in brief.sections)
    assert brief.screenshot_requirements
    table_criteria = [item for item in brief.marking_criteria if "Grafana dashboard" in item.description]
    assert table_criteria
    assert any(span.row_index == 1 for span in table_criteria[0].source_spans)
    assert any(span.column_index is not None for span in table_criteria[0].source_spans)


def test_known_big_data_docx_extracts_exactly_three_assignments() -> None:
    path = Path("assignment_inputs/BigData_Assignments_v5 (2).docx")
    if not path.exists():
        pytest.skip("repository assignment fixture is not present")
    brief = extract_assignment_brief(parse_assignment_document(path))
    assert len(brief.sections) == 3
    assert [section.section_id for section in brief.sections] == ["assignment-1", "assignment-2", "assignment-3"]
    assert all(section.tasks for section in brief.sections)


def test_document_limits_fail_closed(tmp_path: Path) -> None:
    path = tmp_path / "oversized.txt"
    path.write_text("one\ntwo\nthree\n", encoding="utf-8")
    with pytest.raises(DocumentLimitExceeded):
        parse_assignment_document(path, limits=AssignmentDocumentLimits(max_blocks=2))


def test_household_semantic_mapping_is_deterministic_and_placeholder_free() -> None:
    profile = _household_profile()
    first = map_dataset_columns(profile)
    second = map_dataset_columns(profile)
    assert first.model_dump(mode="json") == second.model_dump(mode="json")
    assert first.timestamp_column.column == "event_timestamp"
    assert first.category_grouping_column.column == "event_hour"
    assert first.unresolved_requirements == []
    assert first.placeholders_used is False
    derived = {item.name: item for item in first.semantic_mapping.derived_columns}
    assert derived["event_timestamp"].source_columns == ("Date", "Time")
    assert {"event_hour", "event_weekday"} <= set(derived)

    generated = generate_code_blueprints(2, dataset_profile=profile)
    assert all(not blueprint.placeholders for blueprint in generated.blueprints)
    text = "\n".join(item.generated_content for item in generated.blueprints)
    assert 'TIMESTAMP_COLUMN = "event_timestamp"' in text
    assert 'CATEGORY_COLUMN = "event_hour"' in text
    assert 'TIMESTAMP_COLUMN = "TIMESTAMP_COLUMN"' not in text
    assert 'CATEGORY_COLUMN = "CATEGORY_COLUMN"' not in text


def test_numeric_binning_and_invalid_derived_sources_are_explicit() -> None:
    profile = _household_profile()
    profile = profile.model_copy(update={
        "columns": ["Voltage"], "column_count": 1,
        "detected_date_columns": [], "detected_numeric_columns": ["Voltage"],
    })
    mapping = map_dataset_columns(profile)
    assert any(item.expression_type == "numeric_bin" for item in mapping.semantic_mapping.derived_columns)
    assert mapping.unresolved_requirements
    with pytest.raises(ValueError, match="unknown source"):
        DatasetSemanticMapping(
            source_columns=(SourceColumn(name="Voltage", inferred_type="numeric"),),
            derived_columns=(DerivedColumnPlan(
                name="bad_band", expression_type="numeric_bin", source_columns=("Missing",),
                deterministic_operation="bucket at zero", output_type="string",
                rationale="test", provenance=(),
            ),),
            time_dimension=None, numeric_measures=(), categorical_dimensions=(),
        )


def test_manifest_covers_files_beyond_legacy_position_160_and_detects_changes(tmp_path: Path) -> None:
    project = tmp_path / "many"
    project.mkdir()
    for index in range(170):
        (project / f"module_{index:03}.py").write_text(f"VALUE = {index}\n", encoding="utf-8")
    first = build_project_state_manifest(project, workspace_id="workspace")
    paths = {item.normalized_relative_path for item in first.entries}
    assert len(first.entries) == 170
    assert "module_160.py" in paths and "module_169.py" in paths
    assert build_project_state_manifest(project, workspace_id="workspace").manifest_hash == first.manifest_hash

    (project / "module_160.py").write_text("VALUE = 'changed'\n", encoding="utf-8")
    second = build_project_state_manifest(project, workspace_id="workspace")
    assert second.manifest_hash != first.manifest_hash
    with pytest.raises(ValueError):
        assert_manifest_fresh(first, project)
    (project / "module_169.py").rename(project / "renamed.py")
    third = build_project_state_manifest(project, workspace_id="workspace")
    assert third.manifest_hash != second.manifest_hash
    (project / "added.py").write_text("VALUE = 'added'\n", encoding="utf-8")
    fourth = build_project_state_manifest(project, workspace_id="workspace")
    assert fourth.manifest_hash != third.manifest_hash
    (project / "module_000.py").unlink()
    assert build_project_state_manifest(project, workspace_id="workspace").manifest_hash != fourth.manifest_hash


def test_manifest_exclusions_are_safe_stable_and_fail_closed_on_limits(tmp_path: Path) -> None:
    project = tmp_path / "policy"
    project.mkdir()
    (project / "app.py").write_text("VALUE = 1\n", encoding="utf-8")
    (project / ".env").write_text("SECRET=do-not-hash\n", encoding="utf-8")
    cache = project / ".pytest_cache"
    cache.mkdir()
    (cache / "state").write_text("one", encoding="utf-8")
    first = build_project_state_manifest(project, workspace_id="workspace")
    assert ".env" not in {entry.normalized_relative_path for entry in first.entries}
    assert first.excluded_summary["sensitive_file"] == 1
    try:
        (project / "linked.py").symlink_to(project / "app.py")
    except OSError:
        pass
    else:
        symlink_manifest = build_project_state_manifest(project, workspace_id="workspace")
        assert "linked.py" not in {entry.normalized_relative_path for entry in symlink_manifest.entries}
        assert symlink_manifest.excluded_summary["symlink_file"] == 1
        (project / "linked.py").unlink()
    (project / ".env").write_text("SECRET=changed\n", encoding="utf-8")
    (cache / "state").write_text("two", encoding="utf-8")
    assert build_project_state_manifest(project, workspace_id="workspace").manifest_hash == first.manifest_hash

    for index in range(10):
        (project / f"extra_{index}.py").write_text("pass\n", encoding="utf-8")
    limited = ProjectManifestLimits(max_files=5)
    partial = build_project_state_manifest(project, workspace_id="workspace", limits=limited, require_complete=False)
    assert partial.complete is False
    with pytest.raises(IncompleteProjectManifestError):
        build_project_state_manifest(project, workspace_id="workspace", limits=limited)


def test_plan_definition_is_immutable_while_execution_state_changes(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    plan_hash = job["plan_revision"]["content_hash"]
    revision = copy.deepcopy(job["plan_revision"])
    approved = approve_plan(job, plan_hash=plan_hash)
    active = activate_next_work_unit(approved, root=project)
    assert active["plan_revision"] == revision
    assert active["plan_revision"]["content_hash"] == plan_hash
    assert active["work_unit_execution_states"][0]["status"] == "active"

    tampered = copy.deepcopy(approved)
    tampered["plan"]["work_units"][0]["objective"] = "A different objective"
    with pytest.raises(ProjectDeliveryError) as error:
        activate_next_work_unit(tampered, root=project)
    assert error.value.code == "stale_plan_approval"


@pytest.mark.parametrize("field,replacement", [
    ("expected_files", ["app.py", "new.py"]),
    ("dependencies", ["wu-99"]),
    ("criterion_references", ["criterion-02"]),
])
def test_changed_plan_boundaries_invalidate_approval(tmp_path: Path, field: str, replacement: list[str]) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    approved = approve_plan(job, plan_hash=job["plan_revision"]["content_hash"])
    approved["plan"]["work_units"][0][field] = replacement
    with pytest.raises(ProjectDeliveryError) as error:
        activate_next_work_unit(approved, root=project)
    assert error.value.code == "stale_plan_approval"


def test_incomplete_manifest_cannot_authorize_plan_approval(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    fresh_job = copy.deepcopy(job)
    job["project_state_manifest"]["complete"] = False
    with pytest.raises(ProjectDeliveryError) as error:
        approve_plan(job, plan_hash=job["plan_revision"]["content_hash"])
    assert error.value.code == "incomplete_project_manifest"
    with (project / "oversized.py").open("wb") as handle:
        handle.truncate(10 * 1024 * 1024 + 1)
    criterion = fresh_job["specification"]["acceptance_criteria"][0]
    with pytest.raises(IncompleteProjectManifestError):
        run_deterministic_verifier(
            fresh_job, root=project, criterion_id=criterion["criterion_id"]
        )


def test_revision_id_change_cannot_reuse_an_approval(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    approved = approve_plan(job, plan_hash=job["plan_revision"]["content_hash"])
    approved["plan_revision"]["plan_revision_id"] = "fraudulent-revision"
    with pytest.raises(ProjectDeliveryError) as error:
        activate_next_work_unit(approved, root=project)
    assert error.value.code == "approval_required"


def test_legacy_approval_cannot_authorize_new_plan_revision(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    legacy = copy.deepcopy(job)
    legacy.pop("plan_revision", None)
    legacy["plan_approval"] = {"plan_hash": legacy["plan"]["plan_hash"]}
    with pytest.raises(ProjectDeliveryError) as error:
        approve_plan(legacy, plan_hash=legacy["plan"]["plan_hash"])
    assert error.value.code == "migration_reapproval_required"


def test_verifier_result_is_typed_hash_bound_and_stale_after_change(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    criterion = job["specification"]["acceptance_criteria"][0]
    mode = VerificationMode(criterion["verification_mode"])
    result = run_deterministic_verifier(job, root=project, criterion_id=criterion["criterion_id"])
    assert result.outcome.value == "passed"
    assert result.performed_checks and all(check.outcome == "passed" for check in result.performed_checks)
    assert_verifier_result_fresh(result, job, root=project)

    with pytest.raises(ProjectDeliveryError) as missing:
        record_verification(
            job, work_unit_id="wu-01", criterion_id=criterion["criterion_id"],
            state=VerificationState.SATISFIED, method=mode,
            evidence_references=["analysis-id"], structural_analysis_references=[job["analysis_id"]],
        )
    assert missing.value.code == "missing_checker"

    tampered = result.model_copy(update={"summary": "claimed pass"})
    with pytest.raises(ProjectVerifierError):
        assert_verifier_result_fresh(tampered, job, root=project)
    superseded = copy.deepcopy(job)
    superseded["plan_revision"]["plan_revision_id"] = "new-revision"
    with pytest.raises(ProjectVerifierError):
        assert_verifier_result_fresh(result, superseded, root=project)
    (project / "unrelated.py").write_text("VALUE = 1\n", encoding="utf-8")
    with pytest.raises(ProjectVerifierError) as stale:
        assert_verifier_result_fresh(result, job, root=project)
    assert stale.value.code == "stale_verifier_result"


def test_file_ast_configuration_exact_and_manual_verifiers_compute_results(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    (project / "settings.json").write_text('{"feature":{"enabled":true}}\n', encoding="utf-8")
    base = _delivery_job(project)
    criterion_id = base["specification"]["acceptance_criteria"][0]["criterion_id"]

    file_job = copy.deepcopy(base)
    criterion = file_job["specification"]["acceptance_criteria"][0]
    criterion.update({"verification_mode": "file_presence", "checker_rule": {
        "rule_type": "file_exists", "relative_path": "app.py", "selector": None, "expected_value": None,
    }})
    assert run_deterministic_verifier(file_job, root=project, criterion_id=criterion_id).outcome.value == "passed"

    config_job = copy.deepcopy(base)
    criterion = config_job["specification"]["acceptance_criteria"][0]
    criterion.update({"verification_mode": "configuration", "checker_rule": {
        "rule_type": "json_value", "relative_path": "settings.json",
        "selector": "feature.enabled", "expected_value": True,
    }})
    assert run_deterministic_verifier(config_job, root=project, criterion_id=criterion_id).outcome.value == "passed"
    criterion["checker_rule"]["expected_value"] = False
    assert run_deterministic_verifier(config_job, root=project, criterion_id=criterion_id).outcome.value == "failed"

    exact_job = copy.deepcopy(base)
    exact_job["specification"]["acceptance_criteria"][0].update({
        "verification_mode": "exact_diff_or_content_assertion", "checker_rule": {
            "rule_type": "text_contains", "relative_path": "app.py",
            "selector": "def greet", "expected_value": None,
        },
    })
    assert run_deterministic_verifier(exact_job, root=project, criterion_id=criterion_id).outcome.value == "passed"

    manual_job = copy.deepcopy(base)
    manual_job["specification"]["acceptance_criteria"][0].update({
        "verification_mode": "manual_user_verification_required", "checker_rule": None,
    })
    assert run_deterministic_verifier(manual_job, root=project, criterion_id=criterion_id).outcome.value == "manual_required"

    (project / "app.py").write_text("def broken(:\n", encoding="utf-8")
    structural = run_deterministic_verifier(base, root=project, criterion_id=criterion_id)
    assert structural.outcome.value == "failed"


def test_changed_criterion_invalidates_verifier_result(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    criterion = job["specification"]["acceptance_criteria"][0]
    result = run_deterministic_verifier(job, root=project, criterion_id=criterion["criterion_id"])
    changed = copy.deepcopy(job)
    changed["specification"]["acceptance_criteria"][0]["requirement"] = "A changed requirement"
    with pytest.raises(ProjectVerifierError) as error:
        assert_verifier_result_fresh(result, changed, root=project)
    assert error.value.code == "stale_verifier_result"


def test_handoff_rejects_project_changes_after_fresh_verification(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    job = _delivery_job(project)
    for criterion in job["specification"]["acceptance_criteria"]:
        mode = VerificationMode(criterion["verification_mode"])
        result = run_deterministic_verifier(job, root=project, criterion_id=criterion["criterion_id"])
        job = record_verification(
            job, work_unit_id=str(result.work_unit_id), criterion_id=criterion["criterion_id"],
            state=VerificationState.SATISFIED, method=mode,
            evidence_references=[result.verifier_result_id],
            relevant_file_hashes={"app.py": hashlib.sha256((project / "app.py").read_bytes()).hexdigest()},
            structural_analysis_references=[job["analysis_id"]] if mode == VerificationMode.STRUCTURAL else [],
            verifier_result=result,
        )
    assert generate_handoff(job, root=project)["handoff"]["completion_status"] == "completed"
    (project / "app.py").write_text("def greet(name): return name\n", encoding="utf-8")
    with pytest.raises(ProjectDeliveryError) as stale:
        generate_handoff(job, root=project)
    assert stale.value.code == "stale_verifier_result"


def test_normalized_stage0_records_survive_repository_reload(tmp_path: Path) -> None:
    project = _delivery_project(tmp_path)
    repository = AnalysisRepository(tmp_path / "stage0.db")
    repository.initialize()
    repository.initialize()
    job = _delivery_job(project)
    criterion = job["specification"]["acceptance_criteria"][0]
    result = run_deterministic_verifier(job, root=project, criterion_id=criterion["criterion_id"])
    job = record_verification(
        job, work_unit_id=str(result.work_unit_id), criterion_id=criterion["criterion_id"],
        state=VerificationState.SATISFIED,
        method=VerificationMode(criterion["verification_mode"]),
        evidence_references=[result.verifier_result_id],
        relevant_file_hashes={"app.py": hashlib.sha256((project / "app.py").read_bytes()).hexdigest()},
        structural_analysis_references=[job["analysis_id"]], verifier_result=result,
    )
    repository.store_project_delivery_job(job)
    repository.store_stage0_audit_event({
        "event_id": "audit-1", "domain": "document", "aggregate_id": "doc-1",
        "operation": "document_parse_completion", "status": "completed",
        "metadata": {"block_count": 4}, "created_at": "2026-01-01T00:00:00+00:00",
    })
    reloaded = AnalysisRepository(tmp_path / "stage0.db").get_project_delivery_job(job["delivery_job_id"])
    assert reloaded["plan_revision"]["content_hash"] == job["plan_revision"]["content_hash"]
    assert reloaded["project_state_manifest"]["manifest_hash"] == job["project_state_manifest"]["manifest_hash"]
    assert reloaded["verifier_results"][0]["result_hash"] == result.result_hash
    with sqlite3.connect(tmp_path / "stage0.db") as connection:
        assert connection.execute("SELECT COUNT(*) FROM project_delivery_plan_revisions").fetchone()[0] == 1
        assert connection.execute("SELECT COUNT(*) FROM project_state_manifests").fetchone()[0] == 1
        assert connection.execute("SELECT COUNT(*) FROM project_delivery_work_unit_states").fetchone()[0] >= 1
        assert connection.execute("SELECT COUNT(*) FROM project_verifier_results").fetchone()[0] == 1
    assert AnalysisRepository(tmp_path / "stage0.db").list_stage0_audit_events("document", "doc-1")[0]["operation"] == "document_parse_completion"
