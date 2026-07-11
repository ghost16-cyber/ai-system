from __future__ import annotations

import hashlib
import json
import os
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

from pydantic import ValidationError

from backend.app.assignments.schemas import (
    AssignmentVerificationSnapshot,
    GroundedAssignmentReport,
    GroundedReportContentBlock,
    GroundedReportSection,
    ReportExportReadiness,
    ReportExportRecord,
    ReportRevision,
)
from backend.app.assignments.verification import load_verification_snapshot


REPORT_SCHEMA_VERSION = 1
EXPORT_SCHEMA_VERSION = 1
MAX_PACKAGE_FILE_BYTES = 20 * 1024 * 1024
SAFE_PACKAGE_EXTENSIONS = {
    ".py", ".md", ".txt", ".csv", ".tsv", ".json", ".ipynb", ".png",
    ".jpg", ".jpeg", ".gif", ".webp", ".pdf", ".docx", ".yml", ".yaml",
    ".toml", ".ini", ".html",
}
SECRET_RE = re.compile(r"(?i)\b(password|passwd|token|secret|api[_-]?key)\b(\s*[:=]\s*)([^\s,;]+)")
KEY_RE = re.compile(r"\bsk-[A-Za-z0-9_-]{12,}\b")


class ReportAssemblyError(ValueError):
    pass


def create_grounded_report(*, report_root, verification_root, assignment_id, workspace_relative, title=None):
    snapshot = load_verification_snapshot(verification_root, assignment_id, workspace_relative)
    now = datetime.now(timezone.utc)
    revision = ReportRevision(revision_id=f"revision-{uuid4().hex}", timestamp=now, user_authored_fields={"title": bool(title)})
    report = GroundedAssignmentReport(
        schema_version=REPORT_SCHEMA_VERSION,
        report_id=f"report-{uuid4().hex}",
        assignment_id=assignment_id,
        workspace=workspace_relative,
        title=_redact((title or f"{assignment_id} Evidence-Grounded Report").strip()),
        created_at=now,
        updated_at=now,
        source_verification_snapshot=snapshot.snapshot_id,
        report_sections=_assemble_sections(snapshot),
        requirement_coverage=_coverage(snapshot),
        evidence_references=_usable_evidence(snapshot),
        unresolved_items=_unresolved(snapshot),
        warnings=[
            "Content is assembled only from recorded requirements and evidence.",
            "Visible placeholders require user-authored completion.",
        ],
        current_revision_id=revision.revision_id,
        revisions=[revision],
        recommended_submission_files=_recommended_files(snapshot),
    )
    _write_report(report_root, report)
    return report


def assemble_grounded_report(*, report_root, verification_root, assignment_id, workspace_relative, report_id):
    current = load_grounded_report(report_root, assignment_id, workspace_relative, report_id)
    snapshot = load_verification_snapshot(verification_root, assignment_id, workspace_relative)
    rebuilt = _assemble_sections(snapshot)
    previous = {section.section_id: section for section in current.report_sections}
    for index, section in enumerate(rebuilt):
        old = previous.get(section.section_id)
        if old:
            rebuilt[index] = section.model_copy(update={
                "title": old.title,
                "user_editable_notes": old.user_editable_notes,
                "inclusion_status": old.inclusion_status,
                "selected_evidence": [ref for ref in old.selected_evidence if ref in section.linked_evidence],
                "grounded_content_blocks": [*section.grounded_content_blocks, *[block for block in old.grounded_content_blocks if block.user_authored]],
            })
    order = {section.section_id: index for index, section in enumerate(current.report_sections)}
    rebuilt.sort(key=lambda item: order.get(item.section_id, len(order)))
    now = datetime.now(timezone.utc)
    revision = ReportRevision(
        revision_id=f"revision-{uuid4().hex}", timestamp=now,
        changed_sections=[item.section_id for item in rebuilt],
        previous_revision_reference=current.current_revision_id,
        user_authored_fields={"action": "assemble_from_latest_snapshot"},
    )
    report = current.model_copy(update={
        "updated_at": now,
        "source_verification_snapshot": snapshot.snapshot_id,
        "report_sections": rebuilt,
        "requirement_coverage": _coverage(snapshot),
        "evidence_references": _usable_evidence(snapshot),
        "unresolved_items": _unresolved(snapshot),
        "current_revision_id": revision.revision_id,
        "revisions": [*current.revisions, revision],
        "recommended_submission_files": _recommended_files(snapshot),
    })
    _write_report(report_root, report)
    return report


def update_grounded_report(*, report_root, assignment_id, workspace_relative, report_id, changes):
    report = load_grounded_report(report_root, assignment_id, workspace_relative, report_id)
    sections = {item.section_id: item for item in report.report_sections}
    changed: list[str] = []
    user_fields: dict[str, Any] = {}
    if "title" in changes:
        title = _redact(str(changes["title"]).strip())
        if not title:
            raise ReportAssemblyError("Report title cannot be empty.")
        report = report.model_copy(update={"title": title})
        user_fields["title"] = title
    for edit in changes.get("sections", []):
        if not isinstance(edit, dict) or edit.get("section_id") not in sections:
            raise ReportAssemblyError("Report section edit references an unknown section.")
        section = sections[edit["section_id"]]
        patch: dict[str, Any] = {}
        if "title" in edit:
            patch["title"] = _redact(str(edit["title"]).strip())
        if "user_editable_notes" in edit:
            patch["user_editable_notes"] = _redact(str(edit["user_editable_notes"]))
        if "inclusion_status" in edit:
            if edit["inclusion_status"] not in {"included", "excluded"}:
                raise ReportAssemblyError("Section inclusion status is invalid.")
            if section.mandatory and edit["inclusion_status"] == "excluded":
                raise ReportAssemblyError("Mandatory report sections cannot be excluded.")
            patch["inclusion_status"] = edit["inclusion_status"]
        if "selected_evidence" in edit:
            selected = edit["selected_evidence"]
            if not isinstance(selected, list) or not set(selected).issubset(set(section.linked_evidence)):
                raise ReportAssemblyError("Selected section evidence must be a subset of valid evidence candidates.")
            patch["selected_evidence"] = selected
            patch["citations"] = selected
        replacements = edit.get("placeholder_replacements", {})
        if replacements:
            if not isinstance(replacements, dict):
                raise ReportAssemblyError("Placeholder replacements must be an object.")
            placeholders = list(section.placeholders)
            blocks = list(section.grounded_content_blocks)
            for placeholder, value in replacements.items():
                if placeholder not in placeholders:
                    raise ReportAssemblyError("Placeholder replacement does not match this section.")
                text = _redact(str(value).strip())
                if not text:
                    raise ReportAssemblyError("Placeholder replacement cannot be empty.")
                placeholders.remove(placeholder)
                blocks.append(_block("user_authored", text, [], True))
            patch.update({"placeholders": placeholders, "grounded_content_blocks": blocks})
        sections[section.section_id] = section.model_copy(update=patch)
        changed.append(section.section_id)
        user_fields[section.section_id] = patch
    order = changes.get("section_order", [item.section_id for item in report.report_sections])
    if not isinstance(order, list) or set(order) != set(sections) or len(order) != len(sections):
        raise ReportAssemblyError("Section order must contain every eligible section exactly once.")
    ordered = [sections[section_id] for section_id in order]
    now = datetime.now(timezone.utc)
    revision = ReportRevision(
        revision_id=f"revision-{uuid4().hex}", timestamp=now,
        changed_sections=list(dict.fromkeys([*changed, *([] if "section_order" not in changes else order)])),
        previous_revision_reference=report.current_revision_id,
        user_authored_fields=user_fields,
    )
    report = report.model_copy(update={
        "report_sections": ordered, "updated_at": now,
        "current_revision_id": revision.revision_id,
        "revisions": [*report.revisions, revision],
        "unresolved_items": [placeholder for section in ordered for placeholder in section.placeholders],
    })
    _write_report(report_root, report)
    return report


def report_export_readiness(report):
    included = [item for item in report.report_sections if item.inclusion_status == "included"]
    placeholders = sum(len(item.placeholders) for item in included)
    stale = sum(item.verification_state == "stale" for item in included)
    unsupported = sum(item.verification_state in {"unsupported", "missing"} for item in included)
    manual = sum(item.verification_state == "requires_manual_review" for item in included)
    failed = sum(any("failed" in warning.lower() or "conflict" in warning.lower() for warning in item.warnings) for item in included)
    missing = sum(item.mandatory and item.verification_state in {"unsupported", "missing"} for item in included)
    blockers = []
    if placeholders: blockers.append(f"{placeholders} unresolved placeholder(s).")
    if stale: blockers.append(f"{stale} section(s) use stale evidence.")
    if missing: blockers.append(f"{missing} mandatory section(s) are unsupported or missing.")
    if manual: blockers.append(f"{manual} section(s) require manual review.")
    if failed: blockers.append(f"{failed} section(s) retain failed or conflicting evidence.")
    covered = sum(bool(item.get("evidence_references")) for item in report.requirement_coverage)
    total = len(report.requirement_coverage)
    return ReportExportReadiness(
        report_id=report.report_id, calculated_at=datetime.now(timezone.utc),
        supported_section_count=sum(item.verification_state in {"verified", "manually_accepted", "partially_supported"} for item in included),
        unsupported_section_count=unsupported,
        unresolved_placeholder_count=placeholders,
        stale_evidence_count=stale,
        failed_evidence_count=failed,
        manual_review_count=manual,
        missing_mandatory_section_count=missing,
        traceability_coverage_percentage=round((covered / total * 100) if total else 0.0, 1),
        export_blockers=blockers,
        warnings=["Export readiness is separate from assignment readiness and requires final human review."],
        status="eligible_for_final_human_submission_review" if not blockers else "blocked",
    )


def _assemble_sections(snapshot):
    sections = [
        GroundedReportSection(
            section_id="title-page", title="Title Page Metadata",
            purpose="Identify the bound assignment and verification source.",
            grounded_content_blocks=[
                _block("grounded_statement", f"Assignment ID: {snapshot.assignment_id}", []),
                _block("evidence_note", f"Verification snapshot: {snapshot.snapshot_id}", []),
            ], verification_state="verified", mandatory=True,
        ),
        GroundedReportSection(
            section_id="objectives", title="Assignment Objectives and Requirement Interpretation",
            purpose="Present requirements recorded by the assignment workflow.",
            originating_requirement_ids=[item.requirement_id for item in snapshot.requirements],
            grounded_content_blocks=[_block("requirement", f"Recorded objective: {item.title}. {item.description}", []) for item in snapshot.requirements],
            verification_state="verified", mandatory=True,
        ),
    ]
    groups = [
        ("architecture", "System Architecture or Workflow", ("architecture", "workflow", "pipeline")),
        ("implementation", "Implementation Description", ("code", "implementation", "python", "notebook", "configuration")),
        ("dataset", "Dataset Description", ("dataset", "csv", "tsv")),
        ("testing", "Testing and Validation", ("test", "validation")),
        ("evidence", "Evidence and Screenshots", ("screenshot", "dashboard", "terminal_output", "image")),
        ("discussion", "Discussion and Report Evidence", ("report", "appendix", "report_answer")),
    ]
    for section_id, title, terms in groups:
        requirements = [item for item in snapshot.requirements if any(term in f"{item.requirement_category} {item.required_deliverable_type}" for term in terms)]
        if requirements:
            sections.append(_requirements_section(section_id, title, requirements, snapshot))
    execution_requirements = [item for item in snapshot.requirements if item.linked_execution_evidence]
    if execution_requirements:
        sections.append(_requirements_section("execution-results", "Execution Results", execution_requirements, snapshot))
    all_warnings = list(dict.fromkeys(warning for item in snapshot.requirements for warning in item.warnings))
    limitation_placeholder = "[Required: add user-authored limitations grounded in observed evidence.]"
    sections.append(GroundedReportSection(
        section_id="limitations", title="Limitations",
        purpose="Retain evidence limitations and verification constraints.",
        grounded_content_blocks=[_block("evidence_note", warning, []) for warning in all_warnings] or [_block("placeholder", limitation_placeholder, [])],
        verification_state="partially_supported" if all_warnings else "missing",
        placeholders=[] if all_warnings else [limitation_placeholder], mandatory=True,
    ))
    usable = _usable_evidence(snapshot)
    coverage_by_id = {item["requirement_id"]: item for item in _coverage(snapshot)}
    sections.append(GroundedReportSection(
        section_id="traceability", title="Requirement-to-Evidence Traceability Matrix",
        purpose="Map every requirement to usable evidence and unresolved status.",
        originating_requirement_ids=[item.requirement_id for item in snapshot.requirements],
        grounded_content_blocks=[_block("evidence_note", f"{item.requirement_id}: {item.status}; evidence: {', '.join(coverage_by_id[item.requirement_id]['evidence_references']) or 'none'}.", coverage_by_id[item.requirement_id]["evidence_references"]) for item in snapshot.requirements],
        linked_evidence=usable, selected_evidence=usable, citations=usable,
        verification_state="verified" if all(item.status == "verified" for item in snapshot.requirements) else "partially_supported",
        mandatory=True,
    ))
    files = list(dict.fromkeys(path for item in snapshot.requirements for path in item.linked_workspace_files))
    if files:
        sections.append(GroundedReportSection(
            section_id="appendix", title="Appendix: Selected Deliverables",
            purpose="List inventoried deliverables without asserting their behavior.",
            grounded_content_blocks=[_block("evidence_note", f"Inventoried workspace file: {path}", [f"file:{path}"]) for path in files],
            linked_evidence=[f"file:{path}" for path in files], selected_evidence=[f"file:{path}" for path in files], citations=[f"file:{path}" for path in files],
            verification_state="partially_supported", mandatory=False,
        ))
    return sections


def _requirements_section(section_id, title, requirements, snapshot):
    latest = _latest_reviews(snapshot)
    rejected = {ref for ref, review in latest.items() if review.decision == "rejected"}
    blocks = []; evidence = []; warnings = []; placeholders = []
    for item in requirements:
        refs = [ref for ref in _requirement_refs(item) if ref not in rejected]
        evidence.extend(refs); warnings.extend(item.warnings)
        if item.status == "verified":
            manual = any(latest.get(ref) and latest[ref].decision == "accepted" for ref in refs)
            text = f"Evidence for requirement '{item.title}' was explicitly accepted during manual review." if manual else f"Recorded evidence supports requirement '{item.title}'."
            if any(ref.startswith("assignment-command:") for ref in refs) and "test" in f"{item.requirement_category} {item.required_deliverable_type}":
                text = f"Fresh controlled test evidence supports requirement '{item.title}'."
            blocks.append(_block("grounded_statement", text, refs))
        elif item.status == "detected":
            blocks.append(_block("evidence_note", f"Candidate files exist for '{item.title}': {', '.join(item.linked_workspace_files)}. Functional correctness remains unverified.", refs))
        elif item.status in {"partially_verified", "requires_manual_review"}:
            blocks.append(_block("evidence_note", f"Requirement '{item.title}' is partially supported or still requires review.", refs))
            placeholder = "[Required: add a reviewed explanation supported by the linked evidence.]"
            placeholders.append(placeholder); blocks.append(_block("placeholder", placeholder, []))
        elif item.status == "failed":
            blocks.append(_block("evidence_note", f"Recorded validation or review evidence for '{item.title}' indicates failure.", refs))
            placeholder = "[Required: resolve or explain the failed validation without changing the audit record.]"
            placeholders.append(placeholder); blocks.append(_block("placeholder", placeholder, []))
        elif item.status == "missing":
            placeholder = f"[Required: provide evidence and user-authored content for '{item.title}'.]"
            placeholders.append(placeholder); blocks.append(_block("placeholder", placeholder, []))
    if rejected:
        warnings.append("Rejected evidence remains in the verification audit but is excluded from report support.")
    return GroundedReportSection(
        section_id=section_id, title=title,
        purpose="Assemble only requirement-grounded statements and visible placeholders.",
        originating_requirement_ids=[item.requirement_id for item in requirements],
        grounded_content_blocks=blocks, linked_evidence=list(dict.fromkeys(evidence)), selected_evidence=list(dict.fromkeys(evidence)),
        citations=list(dict.fromkeys(evidence)), verification_state=_section_state(requirements, latest),
        warnings=list(dict.fromkeys(warnings)), placeholders=list(dict.fromkeys(placeholders)), mandatory=True,
    )


def _section_state(requirements, latest):
    statuses = {item.status for item in requirements}
    warnings = " ".join(warning.lower() for item in requirements for warning in item.warnings)
    if "stale" in warnings: return "stale"
    if "conflict" in warnings: return "requires_manual_review"
    if "failed" in statuses: return "unsupported"
    if "missing" in statuses: return "missing"
    if "requires_manual_review" in statuses: return "requires_manual_review"
    if statuses & {"partially_verified", "detected"}: return "partially_supported"
    if statuses == {"not_applicable"}: return "not_applicable"
    accepted = {ref for ref, review in latest.items() if review.decision == "accepted"}
    if accepted and all(any(ref in accepted for ref in _requirement_refs(item)) for item in requirements): return "manually_accepted"
    return "verified"


def _latest_reviews(snapshot):
    latest = {}
    for review in sorted(snapshot.manual_reviews, key=lambda item: item.timestamp):
        latest[review.evidence_reference] = review
    return latest


def _requirement_refs(item):
    return [*[f"file:{path}" for path in item.linked_workspace_files], *item.linked_execution_evidence]


def _coverage(snapshot):
    rejected = {ref for ref, review in _latest_reviews(snapshot).items() if review.decision == "rejected"}
    return [{
        "requirement_id": item.requirement_id,
        "status": item.status,
        "evidence_references": [ref for ref in _requirement_refs(item) if ref not in rejected],
        "warnings": item.warnings,
    } for item in snapshot.requirements]


def _usable_evidence(snapshot):
    rejected = {ref for ref, review in _latest_reviews(snapshot).items() if review.decision == "rejected"}
    return list(dict.fromkeys(ref for item in snapshot.requirements for ref in _requirement_refs(item) if ref not in rejected))


def _recommended_files(snapshot):
    usable = set(_usable_evidence(snapshot))
    return sorted({path for item in snapshot.requirements for path in item.linked_workspace_files if f"file:{path}" in usable})


def _unresolved(snapshot):
    return [item.title for item in snapshot.requirements if item.status in {"missing", "failed", "partially_verified", "requires_manual_review"}]


def _block(kind, text, refs, user=False):
    return GroundedReportContentBlock(block_id=f"block-{uuid4().hex}", block_type=kind, text=_redact(text), evidence_references=refs, user_authored=user)


def export_grounded_report(*, report_root, project_root, assignment_id, workspace_relative, report_id, export_format, selected_files):
    report = load_grounded_report(report_root, assignment_id, workspace_relative, report_id)
    if export_format not in {"markdown", "json", "docx", "zip"}:
        raise ReportAssemblyError("Unsupported report export format.")
    root = Path(project_root).resolve(); workspace = (root / workspace_relative).resolve()
    _require_inside(workspace, root, "Report workspace")
    selected = _validate_selected_files(workspace, selected_files, set(report.recommended_submission_files))
    export_id = f"export-{uuid4().hex}"
    directory = _export_directory(report_root, assignment_id, workspace_relative, report_id, export_id)
    directory.mkdir(parents=True, exist_ok=False)
    readiness = report_export_readiness(report)
    stem = _safe_filename(f"{assignment_id}-{report.title}")
    if export_format == "markdown":
        target = directory / f"{stem}.md"; target.write_text(_report_markdown(report, readiness), encoding="utf-8"); media = "text/markdown"
    elif export_format == "json":
        target = directory / f"{stem}-manifest.json"; target.write_text(json.dumps(_report_manifest(report, readiness, selected), indent=2, sort_keys=True) + "\n", encoding="utf-8"); media = "application/json"
    elif export_format == "docx":
        target = directory / f"{stem}.docx"; _write_docx(report, readiness, target); media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        target = directory / f"{stem}-submission.zip"; _write_zip(report, readiness, target, workspace, selected); media = "application/zip"
    record = ReportExportRecord(
        schema_version=EXPORT_SCHEMA_VERSION, export_id=export_id, report_id=report_id,
        assignment_id=assignment_id, workspace=workspace_relative, format=export_format,
        created_at=datetime.now(timezone.utc), filename=target.name, media_type=media,
        sha256=_sha256(target), size=target.stat().st_size, selected_files=selected,
        readiness=readiness,
        warnings=[] if readiness.status != "blocked" else ["Export contains unresolved readiness blockers."],
    )
    _write_json_atomic(directory / "export-record.json", record.model_dump(mode="json"))
    _write_report(report_root, report.model_copy(update={"export_status": "exported", "updated_at": datetime.now(timezone.utc)}))
    return record


def list_grounded_reports(report_root, assignment_id, workspace_relative):
    directory = _binding_directory(report_root, assignment_id, workspace_relative) / "reports"
    return [] if not directory.is_dir() else [_load_report_path(path, assignment_id, workspace_relative) for path in sorted(directory.glob("report-*.json"))]


def load_grounded_report(report_root, assignment_id, workspace_relative, report_id):
    if not re.fullmatch(r"report-[a-f0-9]{32}", report_id):
        raise ReportAssemblyError("Invalid report id.")
    path = _binding_directory(report_root, assignment_id, workspace_relative) / "reports" / f"{report_id}.json"
    if not path.is_file(): raise FileNotFoundError("Assignment report not found.")
    return _load_report_path(path, assignment_id, workspace_relative)


def list_report_exports(report_root, assignment_id, workspace_relative, report_id):
    load_grounded_report(report_root, assignment_id, workspace_relative, report_id)
    root = _binding_directory(report_root, assignment_id, workspace_relative) / "exports" / report_id
    return [] if not root.is_dir() else [_load_export(path / "export-record.json", assignment_id, workspace_relative, report_id) for path in sorted(root.glob("export-*"))]


def resolve_report_export(report_root, assignment_id, workspace_relative, report_id, export_id):
    if not re.fullmatch(r"export-[a-f0-9]{32}", export_id): raise ReportAssemblyError("Invalid export id.")
    directory = _binding_directory(report_root, assignment_id, workspace_relative) / "exports" / report_id / export_id
    record = _load_export(directory / "export-record.json", assignment_id, workspace_relative, report_id)
    target = (directory / record.filename).resolve(); _require_inside(target, directory.resolve(), "Report export")
    if not target.is_file() or _sha256(target) != record.sha256: raise ReportAssemblyError("Report export failed its integrity check.")
    return record, target


def _report_markdown(report, readiness):
    lines = [f"# {report.title}", "", f"Assignment: {report.assignment_id}", f"Verification snapshot: {report.source_verification_snapshot}", "", "> Evidence-grounded draft. Final human review is required.", ""]
    for section in report.report_sections:
        if section.inclusion_status == "excluded": continue
        lines.extend([f"## {section.title}", "", f"Section state: {section.verification_state}", ""])
        for block in section.grounded_content_blocks:
            lines.extend([("User-authored: " if block.user_authored else "") + block.text, ""])
            selected_refs = [ref for ref in block.evidence_references if ref in section.selected_evidence]
            if selected_refs: lines.extend([f"Evidence: {', '.join(selected_refs)}", ""])
        if section.user_editable_notes: lines.extend([f"User-authored notes: {section.user_editable_notes}", ""])
        for warning in section.warnings: lines.extend([f"> Warning: {warning}", ""])
    lines.extend(["## Export Readiness", "", f"Status: {readiness.status}", ""])
    lines.extend(f"- {blocker}" for blocker in readiness.export_blockers)
    return "\n".join(lines).rstrip() + "\n"


def _report_manifest(report, readiness, selected):
    return {
        "schema_version": REPORT_SCHEMA_VERSION,
        "report": report.model_dump(mode="json"),
        "export_readiness": readiness.model_dump(mode="json"),
        "selected_files": selected,
        "absolute_paths_included": False,
        "academic_correctness_guaranteed": False,
    }


def _write_zip(report, readiness, target, workspace, selected):
    execution = [{"evidence_reference": ref, "redacted_summary": "Controlled execution evidence linked; internal executor state excluded."} for ref in report.evidence_references if ref.startswith("assignment-command:")]
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("report/report.md", _report_markdown(report, readiness))
        archive.writestr("report/traceability-manifest.json", json.dumps(_report_manifest(report, readiness, selected), indent=2, sort_keys=True) + "\n")
        archive.writestr("report/readiness.json", json.dumps(readiness.model_dump(mode="json"), indent=2, sort_keys=True) + "\n")
        archive.writestr("report/redacted-execution-summaries.json", json.dumps(execution, indent=2, sort_keys=True) + "\n")
        for item in selected: archive.write(workspace / item["relative_path"], f"workspace/{item['relative_path']}")


def _write_docx(report, readiness, target):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as error:
        raise ReportAssemblyError("python-docx is required for DOCX export.") from error
    doc = Document(); section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (("Heading 1", 16, 16, 8, "2E74B5"), ("Heading 2", 13, 12, 6, "2E74B5"), ("Heading 3", 12, 8, 4, "1F4D78")):
        style = doc.styles[name]; style.font.name = "Calibri"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color); style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    section.header.paragraphs[0].text = "Evidence-Grounded Assignment Report"
    section.footer.paragraphs[0].text = f"{report.assignment_id} | Human review required"
    cover = doc.add_paragraph(); cover.alignment = WD_ALIGN_PARAGRAPH.CENTER; cover.paragraph_format.space_before = Pt(110); cover.paragraph_format.space_after = Pt(8)
    run = cover.add_run(report.title); run.bold = True; run.font.name = "Calibri"; run.font.size = Pt(28); run.font.color.rgb = RGBColor.from_string("203748")
    subtitle = doc.add_paragraph(f"Evidence-grounded report draft\nAssignment {report.assignment_id}"); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER; subtitle.paragraph_format.space_after = Pt(72)
    meta = doc.add_paragraph(f"Created: {report.created_at.isoformat()}\nSource verification: {report.source_verification_snapshot}\nExport readiness: {readiness.status}"); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    for item in report.report_sections:
        if item.inclusion_status == "excluded" or item.section_id == "title-page": continue
        doc.add_heading(item.title, level=1); status = doc.add_paragraph(); status.add_run(f"Section state: {item.verification_state}").bold = True
        for block in item.grounded_content_blocks:
            p = doc.add_paragraph()
            if block.user_authored: p.add_run("User-authored: ").bold = True
            p.add_run(block.text)
            selected_refs = [ref for ref in block.evidence_references if ref in item.selected_evidence]
            if selected_refs:
                cite = doc.add_paragraph(f"Evidence: {', '.join(selected_refs)}"); cite.paragraph_format.space_before = Pt(4); cite.paragraph_format.space_after = Pt(4)
        if item.user_editable_notes:
            p = doc.add_paragraph(); p.add_run("User-authored notes: ").bold = True; p.add_run(item.user_editable_notes)
        for warning in item.warnings:
            p = doc.add_paragraph(); warning_run = p.add_run(f"Warning: {warning}"); warning_run.font.color.rgb = RGBColor.from_string("9B1C1C")
    doc.add_heading("Requirement Traceability", level=1)
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table.autofit = False
    for cell, text in zip(table.rows[0].cells, ("Requirement", "Status", "Evidence")): cell.text = text
    for row in report.requirement_coverage:
        cells = table.add_row().cells; cells[0].text = str(row["requirement_id"]); cells[1].text = str(row["status"]); cells[2].text = ", ".join(row["evidence_references"]) or "None"
    widths = [2160, 1800, 5400]; tbl = table._tbl; grid = tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    tblW = tbl.tblPr.first_child_found_in("w:tblW")
    if tblW is None: tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360"); tblW.set(qn("w:type"), "dxa")
    if tblW.getparent() is None: tbl.tblPr.append(tblW)
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa"); tbl.tblPr.append(tblInd)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440); tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW(); tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
    doc.save(target)


def _validate_selected_files(workspace, values, allowed_paths):
    if not isinstance(values, list): raise ReportAssemblyError("Selected files must be an explicit list.")
    result = []
    for raw in values:
        path = Path(str(raw))
        if path.is_absolute() or ".." in path.parts: raise ReportAssemblyError("Selected files must be confined relative paths.")
        if path.as_posix() not in allowed_paths: raise ReportAssemblyError("Selected file is not a verified or accepted report deliverable.")
        candidate = workspace / path; current = workspace
        for part in path.parts:
            current = current / part
            if current.is_symlink(): raise ReportAssemblyError("Selected files cannot traverse symlinks.")
        resolved = candidate.resolve(); _require_inside(resolved, workspace, "Selected submission file")
        if not resolved.is_file(): raise ReportAssemblyError(f"Selected file not found: {path.as_posix()}")
        if resolved.suffix.lower() not in SAFE_PACKAGE_EXTENSIONS: raise ReportAssemblyError("Selected file type is not allowed.")
        if resolved.stat().st_size > MAX_PACKAGE_FILE_BYTES: raise ReportAssemblyError("Selected file exceeds the size limit.")
        result.append({"relative_path": path.as_posix(), "sha256": _sha256(resolved), "size": resolved.stat().st_size})
    return result


def _write_report(root, report):
    _write_json_atomic(_binding_directory(root, report.assignment_id, report.workspace) / "reports" / f"{report.report_id}.json", report.model_dump(mode="json"))


def _load_report_path(path, assignment_id, workspace):
    try: report = GroundedAssignmentReport.model_validate(json.loads(path.read_text(encoding="utf-8")))
    except (OSError, json.JSONDecodeError, ValidationError) as error: raise ReportAssemblyError("Assignment report record is corrupt or unsupported.") from error
    if report.schema_version != REPORT_SCHEMA_VERSION: raise ReportAssemblyError("Assignment report schema version is unsupported.")
    if report.assignment_id != assignment_id or report.workspace != workspace: raise ReportAssemblyError("Report belongs to a different assignment or workspace.")
    return report


def _load_export(path, assignment_id, workspace, report_id):
    if not path.is_file(): raise FileNotFoundError("Report export not found.")
    try: record = ReportExportRecord.model_validate(json.loads(path.read_text(encoding="utf-8")))
    except (OSError, json.JSONDecodeError, ValidationError) as error: raise ReportAssemblyError("Report export record is corrupt or unsupported.") from error
    if record.schema_version != EXPORT_SCHEMA_VERSION or record.assignment_id != assignment_id or record.workspace != workspace or record.report_id != report_id: raise ReportAssemblyError("Report export record is unsupported or belongs to another assignment.")
    return record


def _binding_directory(root, assignment_id, workspace):
    digest = hashlib.sha256(f"{assignment_id}|{workspace}".encode()).hexdigest()[:20]
    return Path(root).resolve() / f"{_safe_filename(assignment_id)}-{digest}"


def _export_directory(root, assignment_id, workspace, report_id, export_id):
    return _binding_directory(root, assignment_id, workspace) / "exports" / report_id / export_id


def _safe_filename(value):
    return re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-.")[:100] or "assignment-report"


def _write_json_atomic(path, value):
    path.parent.mkdir(parents=True, exist_ok=True); temporary = path.with_name(f".{path.name}.{uuid4().hex}.tmp")
    temporary.write_text(json.dumps(value, indent=2, sort_keys=True) + "\n", encoding="utf-8"); os.replace(temporary, path)


def _sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""): digest.update(chunk)
    return digest.hexdigest()


def _require_inside(path, root, label):
    try: path.relative_to(root)
    except ValueError as error: raise ReportAssemblyError(f"{label} must stay inside the bound workspace.") from error


def _redact(value):
    return KEY_RE.sub("<redacted-api-key>", SECRET_RE.sub(r"\1\2<redacted>", value))
