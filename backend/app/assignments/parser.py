from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

from backend.app.core.path_utils import normalize_path_for_platform
from backend.app.assignments.schemas import DocumentBlock, ParsedAssignmentDocument, SourceSpan


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx"}
DOCX_MISSING_MESSAGE = "python-docx is required to parse .docx files. Install python-docx or paste the assignment text."


@dataclass(frozen=True, slots=True)
class AssignmentDocumentLimits:
    max_source_bytes: int = 50 * 1024 * 1024
    max_blocks: int = 10_000
    max_tables: int = 100
    max_rows_per_table: int = 2_000
    max_cells: int = 20_000
    max_chars_per_block: int = 50_000
    max_total_chars: int = 2_000_000


DEFAULT_DOCUMENT_LIMITS = AssignmentDocumentLimits()


class AssignmentDocumentError(ValueError):
    code = "unsupported_document_structure"


class DocumentLimitExceeded(AssignmentDocumentError):
    code = "document_limit_exceeded"


def parse_assignment_document(
    path: str | Path, *, limits: AssignmentDocumentLimits = DEFAULT_DOCUMENT_LIMITS,
) -> ParsedAssignmentDocument:
    source = normalize_path_for_platform(path).path
    if not source.exists():
        raise FileNotFoundError(f"Assignment document not found: {source}")
    if not source.is_file():
        raise ValueError("Assignment document path must point to a file, but this path is a folder.")
    suffix = source.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported assignment document extension. Supported extensions: .docx, .md, .txt.")

    if source.stat().st_size > limits.max_source_bytes:
        raise DocumentLimitExceeded(
            f"Document exceeds {limits.max_source_bytes} source bytes."
        )
    warnings: list[str] = []
    source_id = _stream_source_id(source)
    if suffix in {".txt", ".md"}:
        raw_text = source.read_text(encoding="utf-8", errors="ignore")
        blocks = _text_blocks(raw_text, source=source, source_id=source_id, limits=limits)
        text = _compatibility_text(blocks)
    else:
        blocks, docx_warnings = _read_docx(source, source_id=source_id, limits=limits)
        text = _compatibility_text(blocks)
        warnings.extend(docx_warnings)
    text = _normalize_text(text)
    title = _title_from_text(text) or source.stem.replace("_", " ").replace("-", " ").strip()
    return ParsedAssignmentDocument(
        document_id=f"assignment-{source_id[:16]}",
        title=title,
        source_path=str(source),
        extracted_text=text,
        document_blocks=blocks,
        created_at=datetime.now(UTC),
        warnings=warnings,
    )


def _read_docx(
    path: Path, *, source_id: str, limits: AssignmentDocumentLimits,
) -> tuple[list[DocumentBlock], list[str]]:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn
    except ImportError as error:
        raise ValueError(DOCX_MISSING_MESSAGE) from error
    document = Document(str(path))
    blocks: list[DocumentBlock] = []
    paragraph_index = 0
    table_index = 0
    cell_count = 0
    total_chars = 0

    def append_block(
        block_type: str, raw_text: str, *, paragraph_no: int | None = None,
        table_no: int | None = None, row_no: int | None = None,
        column_no: int | None = None, style_name: str | None = None,
        heading_level: int | None = None, table_id: str | None = None,
        metadata: dict | None = None,
    ) -> DocumentBlock | None:
        nonlocal total_chars
        normalized = _normalize_block_text(raw_text)
        if not normalized and block_type not in {"table"}:
            return None
        if len(normalized) > limits.max_chars_per_block:
            raise DocumentLimitExceeded(
                f"Document block exceeds {limits.max_chars_per_block} characters."
            )
        if len(blocks) >= limits.max_blocks:
            raise DocumentLimitExceeded(f"Document exceeds {limits.max_blocks} extracted blocks.")
        if total_chars + len(normalized) > limits.max_total_chars:
            raise DocumentLimitExceeded(
                f"Document exceeds {limits.max_total_chars} extracted characters."
            )
        order_index = len(blocks)
        span = SourceSpan(
            source_id=source_id, source_type="docx", document_path_or_name=str(path),
            block_index=order_index, paragraph_index=paragraph_no,
            table_index=table_no, row_index=row_no, column_index=column_no,
        )
        location = f"p={paragraph_no};t={table_no};r={row_no};c={column_no};type={block_type}"
        block = DocumentBlock(
            block_id=_block_id(source_id, location, raw_text), block_type=block_type,
            order_index=order_index, text=normalized, raw_text=raw_text,
            source_span=span, style_name=style_name, heading_level=heading_level,
            table_id=table_id, row_index=row_no, column_index=column_no,
            metadata=metadata or {},
        )
        blocks.append(block)
        total_chars += len(normalized)
        return block

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            raw = paragraph.text
            style_name = paragraph.style.name if paragraph.style is not None else None
            heading_level = _heading_level(style_name)
            numbered = bool(
                paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
            )
            block_type = "heading" if heading_level else "list_item" if numbered else "paragraph"
            append_block(
                block_type, raw, paragraph_no=paragraph_index, style_name=style_name,
                heading_level=heading_level, metadata={"numbered": numbered},
            )
            paragraph_index += 1
            continue
        if child.tag != qn("w:tbl"):
            continue
        if table_index >= limits.max_tables:
            raise DocumentLimitExceeded(f"Document exceeds {limits.max_tables} tables.")
        table = Table(child, document)
        if len(table.rows) > limits.max_rows_per_table:
            raise DocumentLimitExceeded(
                f"Table {table_index} exceeds {limits.max_rows_per_table} rows."
            )
        table_id = _block_id(source_id, f"table={table_index}", "")
        table_text_rows: list[str] = []
        unique_cells_by_row: list[list[tuple[int, str, int]]] = []
        for row_no, row in enumerate(table.rows):
            seen_tc: set[int] = set()
            row_values: list[tuple[int, str, int]] = []
            for column_no, cell in enumerate(row.cells):
                identity = id(cell._tc)
                if identity in seen_tc:
                    continue
                seen_tc.add(identity)
                raw_cell = "\n".join(p.text for p in cell.paragraphs)
                row_values.append((column_no, raw_cell, identity))
            unique_cells_by_row.append(row_values)
            table_text_rows.append(" | ".join(_normalize_block_text(value) for _, value, _ in row_values))
        append_block(
            "table", "\n".join(table_text_rows), table_no=table_index,
            table_id=table_id, metadata={"row_count": len(table.rows)},
        )
        for row_no, row_values in enumerate(unique_cells_by_row):
            row_text = " | ".join(_normalize_block_text(value) for _, value, _ in row_values)
            row_block = append_block(
                "table_row", row_text, table_no=table_index, row_no=row_no,
                table_id=table_id, metadata={"cell_count": len(row_values)},
            )
            cell_block_ids: list[str] = []
            cell_source_spans: list[dict] = []
            for column_no, raw_cell, _identity in row_values:
                cell_count += 1
                if cell_count > limits.max_cells:
                    raise DocumentLimitExceeded(f"Document exceeds {limits.max_cells} table cells.")
                cell_block = append_block(
                    "table_cell", raw_cell, table_no=table_index, row_no=row_no,
                    column_no=column_no, table_id=table_id,
                )
                if cell_block is not None:
                    cell_block_ids.append(cell_block.block_id)
                    cell_source_spans.append(cell_block.source_span.model_dump(mode="json"))
            if row_block is not None and cell_block_ids:
                blocks[blocks.index(row_block)] = row_block.model_copy(
                    update={"metadata": {**row_block.metadata, "cell_block_ids": cell_block_ids, "cell_source_spans": cell_source_spans}}
                )
        table_index += 1
    if not blocks:
        raise AssignmentDocumentError("The DOCX document did not contain readable paragraphs or tables.")
    return blocks, []


def _text_blocks(
    text: str, *, source: Path, source_id: str, limits: AssignmentDocumentLimits,
) -> list[DocumentBlock]:
    blocks: list[DocumentBlock] = []
    total = 0
    for paragraph_index, raw in enumerate(text.replace("\r\n", "\n").replace("\r", "\n").split("\n")):
        normalized = _normalize_block_text(raw)
        if not normalized:
            continue
        if len(normalized) > limits.max_chars_per_block:
            raise DocumentLimitExceeded(f"Document block exceeds {limits.max_chars_per_block} characters.")
        if len(blocks) >= limits.max_blocks or total + len(normalized) > limits.max_total_chars:
            raise DocumentLimitExceeded("Text document exceeds configured extraction limits.")
        block_type = "heading" if raw.lstrip().startswith("#") else "list_item" if re.match(r"^\s*(?:[-*]|\d+[.)])\s+", raw) else "paragraph"
        order = len(blocks)
        span = SourceSpan(
            source_id=source_id, source_type=source.suffix.lower().lstrip("."),
            document_path_or_name=str(source), block_index=order,
            paragraph_index=paragraph_index,
        )
        heading_level = len(raw) - len(raw.lstrip("#")) if block_type == "heading" else None
        blocks.append(DocumentBlock(
            block_id=_block_id(source_id, f"p={paragraph_index};type={block_type}", raw),
            block_type=block_type, order_index=order, text=normalized.strip("# "),
            raw_text=raw, source_span=span, heading_level=heading_level or None,
            metadata={},
        ))
        total += len(normalized)
    return blocks


def _compatibility_text(blocks: list[DocumentBlock]) -> str:
    visible = [
        block.text for block in blocks
        if block.block_type in {"paragraph", "heading", "list_item", "table_row"} and block.text
    ]
    return "\n".join(visible)


def _normalize_block_text(text: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line)


def _heading_level(style_name: str | None) -> int | None:
    match = re.search(r"(?i)heading\s*([1-9])", style_name or "")
    return int(match.group(1)) if match else None


def _block_id(source_id: str, location: str, raw_text: str) -> str:
    digest = hashlib.sha256(f"{source_id}\0{location}\0{raw_text}".encode("utf-8")).hexdigest()
    return f"doc-block-{digest[:24]}"


def _stream_source_id(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def _normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    normalized = "\n".join(lines)
    while "\n\n\n" in normalized:
        normalized = normalized.replace("\n\n\n", "\n\n")
    return normalized.strip()


def _title_from_text(text: str) -> str | None:
    for line in text.splitlines():
        cleaned = line.strip().strip("#").strip()
        if cleaned:
            return cleaned[:120]
    return None
